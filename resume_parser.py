"""
Enhanced Resume Parser supporting PDF, DOCX, TXT, and Image formats
with improved text extraction and preprocessing.
"""

import os
import re
from typing import Tuple, Optional

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import cv2
except ImportError:
    cv2 = None


class ResumeParser:
    """Enhanced resume parser with robust text extraction."""
    
    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg"]
    
    def extract_text(self, file_path: str, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from resume file.
        
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            ext = os.path.splitext(filename.lower())[1]
            
            if ext not in self.supported_formats:
                return "", f"Unsupported file format: {ext}. Supported: {', '.join(self.supported_formats)}"
            
            # Extract based on file type
            if ext == ".pdf":
                text = self._extract_from_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                text = self._extract_from_docx(file_path)
            elif ext == ".txt":
                text = self._extract_from_txt(file_path)
            elif ext in [".png", ".jpg", ".jpeg"]:
                text = self._extract_from_image(file_path)
            else:
                return "", f"Handler not implemented for {ext}"
            
            # Clean and validate text
            text = self._clean_text(text)
            
            if not text or len(text) < 50:
                return text, "Could not extract sufficient text. File may be corrupted or empty."
            
            return text, None
            
        except Exception as e:
            return "", f"Error extracting text: {str(e)}"
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not pdfplumber:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not docx:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = docx.Document(file_path)
        
        # Extract from paragraphs
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        
        # Extract from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        all_text = "\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n" + "\n".join(tables_text)
        
        return all_text.strip()
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        if not Image or not pytesseract:
            raise ImportError("PIL/pytesseract not installed. Run: pip install Pillow pytesseract")
        
        # Try with preprocessing if opencv is available
        if cv2:
            return self._extract_with_preprocessing(file_path)
        else:
            # Simple extraction
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return text.strip()
    
    def _extract_with_preprocessing(self, file_path: str) -> str:
        """Extract text with image preprocessing for better OCR."""
        # Read image
        img = cv2.imread(file_path)
        
        if img is None:
            raise ValueError("Could not read image file")
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding for better OCR
        _, threshold = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Denoise
        denoised = cv2.fastNlMeansDenoising(threshold, None, 10, 7, 21)
        
        # Save temporarily
        temp_path = file_path + "_processed.png"
        cv2.imwrite(temp_path, denoised)
        
        try:
            # OCR on processed image
            img_pil = Image.open(temp_path)
            text = pytesseract.image_to_string(img_pil, config='--psm 6')
        finally:
            # Cleanup
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
        
        return text.strip()
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def validate_resume(self, text: str) -> Tuple[bool, Optional[str]]:
        """
        Validate if extracted text looks like a resume.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not text or len(text) < 50:
            return False, "Text too short to be a valid resume"
        
        # Check for common resume indicators
        indicators = [
            r'\b(experience|education|skills|projects|work)\b',
            r'\b\d{4}\b',  # Year
            r'@',  # Email
            r'\b(university|college|institute|school)\b',
        ]
        
        matches = sum(1 for pattern in indicators if re.search(pattern, text, re.IGNORECASE))
        
        if matches < 2:
            return False, "Content doesn't appear to be a resume. Please upload a valid resume file."
        
        return True, None

# Singleton instance
resume_parser = ResumeParser()